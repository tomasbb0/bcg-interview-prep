#!/usr/bin/env python3
"""
Generate a tailored Google AE CV as DOCX from the markdown content.
Uses the Feb 2026 DOCX as formatting reference.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '..', 'Google_AE_Application', 'TomásBatalha_Resume_02_2026_GOOGLE_AE_FINAL.docx')

def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 4, "color": "000000"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ('sz', 'val', 'color', 'space'):
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def remove_paragraph_spacing(paragraph):
    """Remove spacing before and after paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(12)

def add_section_header(doc, text):
    """Add a section header with bottom border line."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    # Remove all borders first
    for edge in ('top', 'start', 'end'):
        set_cell_border(cell, **{edge: {"sz": 0, "val": "none", "color": "FFFFFF"}})
    # Bottom border only
    set_cell_border(cell, bottom={"sz": 6, "val": "single", "color": "000000"})
    
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    remove_paragraph_spacing(p)
    
    # Add small spacing after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(2)
    
    return table

def add_job(doc, company_location, role_dates, bullets):
    """Add a job entry with company, role, and bullet points."""
    # Company line
    p_company = doc.add_paragraph()
    run = p_company.add_run(company_location)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    pf = p_company.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(12)
    
    # Role line
    p_role = doc.add_paragraph()
    run = p_role.add_run(role_dates)
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    pf = p_role.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(12)
    
    # Bullets
    for bullet in bullets:
        p = doc.add_paragraph()
        run = p.add_run(f"• {bullet}")
        run.font.size = Pt(8.5)
        run.font.name = 'Calibri'
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.line_spacing = Pt(11)
        pf.left_indent = Inches(0.15)

def add_education(doc, institution, degree, details=None):
    """Add an education entry."""
    p = doc.add_paragraph()
    run = p.add_run(institution)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(12)
    
    p2 = doc.add_paragraph()
    run2 = p2.add_run(degree)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = 'Calibri'
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(1)
    pf2.line_spacing = Pt(12)
    
    if details:
        for detail in details:
            p3 = doc.add_paragraph()
            run3 = p3.add_run(f"• {detail}")
            run3.font.size = Pt(8.5)
            run3.font.name = 'Calibri'
            pf3 = p3.paragraph_format
            pf3.space_before = Pt(0)
            pf3.space_after = Pt(1)
            pf3.line_spacing = Pt(11)
            pf3.left_indent = Inches(0.15)


def create_cv():
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # ==================== HEADER ====================
    # Name
    p_name = doc.add_paragraph()
    run = p_name.add_run("Tomás Maria Burnay Batalha")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 0, 0)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    
    # Contact info
    p_contact = doc.add_paragraph()
    contact_text = "Lisbon, Portugal (Open to Relocation to Dublin) | (+351) 936 124 118 | tomas.b.batalha@gmail.com | linkedin.com/in/tomasmbatalha"
    run = p_contact.add_run(contact_text)
    run.font.size = Pt(8)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(80, 80, 80)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(4)
    
    # ==================== PROFESSIONAL SUMMARY ====================
    add_section_header(doc, "PROFESSIONAL SUMMARY")
    
    p_summary = doc.add_paragraph()
    summary = "Former Google intern and startup Co-Founder with full-cycle new business sales experience. Built a €147K B2B pipeline from zero via multi-channel outreach, consultative selling, and enterprise deal closing. Managed a 2,311-account territory at Amazon, outperforming full-time colleagues. Seeking to bring builder mentality and commercial execution to Google's New Business Sales team."
    run = p_summary.add_run(summary)
    run.font.size = Pt(8.5)
    run.font.name = 'Calibri'
    pf = p_summary.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.line_spacing = Pt(11)
    
    # ==================== WORK EXPERIENCE ====================
    add_section_header(doc, "WORK EXPERIENCE")
    
    # Pairwire
    add_job(doc,
        "Pairwire | New York, USA (Remote)",
        "Co-Founder and Head of Sales | February 2024 – December 2025",
        [
            "Built €147K pipeline of contracted enterprise pilots from zero in 11 months pre-funding through outbound prospecting, cold calling, consulting, and relationship management",
            "Closed 23 enterprise engagements including Fortune 500 and Big 4 firms across 6-9 month full-cycle sales processes",
            "Engineered an AI-powered outbound engine processing 7,000+ leads, automating 80% of prospecting via custom tech stack (Python, Expand.io, AI-enhanced workflows)",
            "Led end-to-end deal management from first touch to contract close, managing entire funnel with no sales team or brand recognition",
            "Created partnership and investor pitch materials; led negotiations with Google, Web Summit, and enterprise decision-makers",
        ])
    
    # Google
    add_job(doc,
        "Google | Lisbon, Portugal",
        "Business Analyst Intern (Large Customer Sales) | May 2024 – August 2024",
        [
            'Led the "Future Trends" segment at Google\'s flagship partner event, presenting personally researched insights on Peak Season activation and media readiness to 237+ senior advertising clients',
            "Designed and automated three globally scalable dashboards using Apps Script and GMP tools to analyze ROAS vs. profit and CPA vs. conversions, projecting $1.41M in annual savings for media spend decisions",
            "Built a reporting platform for Partner Managers using Looker Studio, Analytics 360, and Apps Script, delivering an estimated €315K in annual efficiency gains",
        ])
    
    # Amazon
    add_job(doc,
        "Amazon | Madrid, Spain",
        "Business Development Intern (Amazon Business) | May 2022 – November 2022",
        [
            "Led the B2B retail expansion into Portugal, managing a portfolio of 2,311 companies through prospecting, outreach, and account activation",
            "Outperformed KPI benchmarks by 21x in company account configurations during the Portugal B2B rollout",
            "Raised territory share from 56% below to 25% above team average, outperforming full-time colleagues in the region",
            "Compiled a 19-page strategic report on B2B eCommerce based on 8 company territories and 54 targeted client interviews",
        ])
    
    # EY
    add_job(doc,
        "EY | Lisbon, Portugal",
        "Assurance Associate (Financial Services) | February 2022 – April 2022",
        [
            "Audited asset portfolios worth €41.3 billion across 3 insurance firms, developing rigorous data verification and attention to detail",
        ])
    
    # ==================== EDUCATION ====================
    add_section_header(doc, "EDUCATION")
    
    add_education(doc,
        "Rotterdam School of Management | Rotterdam, Netherlands",
        "MSc in Strategic Management | GPA: 8/10",
        ["Relevant Courses: Applied Machine Learning to Strategy (9.3/10), Python Fundamentals (8.3/10)"])
    
    add_education(doc,
        "National University of Singapore | Singapore",
        "MSc Curriculum Term | MBA-aligned coursework in Strategy and Innovation",
        ["Built a commercial and financial plan for a NUS PhD venture: 70% energy savings, 45% unit margins, 278% ROI in a $1.5B market"])
    
    add_education(doc,
        "Nova School of Business and Economics | Lisbon, Portugal",
        "BSc in Management | GPA: 16/20",
        ["Statistics (20/20), Procurement Management (19/20), Economics of the European Union (19/20)"])
    
    # ==================== LEADERSHIP & ACTIVITIES ====================
    add_section_header(doc, "LEADERSHIP & ACTIVITIES")
    
    leadership_items = [
        ("Nova Thirst Project", "Founder & President: Led 23 volunteers, raised E2,539 for water access, organized 4 events including one with 35 NGOs (2020-2021)"),
        ("Nova Tech Club", "Digital Transformation Consultant: Developed a job platform prototype and organized corporate events and workshops (2020-2021)"),
        ("Nova Social Consulting (WWF)", "Pro Bono Consultant: Assessed organizational structure, identified 155 improvement areas, delivered 6 presentations to senior management (2020-2021)"),
    ]
    
    for name, desc in leadership_items:
        p = doc.add_paragraph()
        run_name = p.add_run(f"{name} — ")
        run_name.bold = True
        run_name.font.size = Pt(8.5)
        run_name.font.name = 'Calibri'
        run_desc = p.add_run(desc)
        run_desc.font.size = Pt(8.5)
        run_desc.font.name = 'Calibri'
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = Pt(11)
    
    # ==================== SKILLS ====================
    add_section_header(doc, "SKILLS")
    
    skills = [
        ("Sales & Marketing:", "Google Ads, YouTube Ads, Google Display Network, CRM (Salesforce), LinkedIn Sales Navigator, Apollo, Full-Cycle Sales, Pipeline Management, Cold Outreach, Consultative Selling"),
        ("Technical:", "Python, Apps Script, Visual Basic for Applications (VBA), SQL, Looker Studio, Power BI, DAX"),
        ("Certifications:", "Google Ads Certified, Bloomberg Market Concepts, Microsoft Excel Specialist, TOEFL iBT"),
        ("Languages:", "Portuguese (Native), English (Fluent), Spanish (Fluent), French (Beginner)"),
    ]
    
    for label, content in skills:
        p = doc.add_paragraph()
        run_label = p.add_run(label + " ")
        run_label.bold = True
        run_label.font.size = Pt(8.5)
        run_label.font.name = 'Calibri'
        run_content = p.add_run(content)
        run_content.font.size = Pt(8.5)
        run_content.font.name = 'Calibri'
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = Pt(11)
    
    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ CV saved to: {OUTPUT_PATH}")
    print(f"   Full path: {os.path.abspath(OUTPUT_PATH)}")
    
    # Also save a copy to Downloads for easy access
    downloads_path = os.path.expanduser('~/Downloads/TomásBatalha_Resume_02_2026_GOOGLE_AE_FINAL.docx')
    doc.save(downloads_path)
    print(f"✅ Also saved to: {downloads_path}")


if __name__ == '__main__':
    create_cv()
