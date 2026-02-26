"""
Generate ATS-optimized CV (table-free) for Google Careers upload.
Uses python-docx to create a clean Word document with proper heading styles,
no tables, no columns, no text boxes — just paragraphs and bullets.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = os.path.expanduser("~/Downloads/TomásBatalha_Resume_GOOGLE_AE_ATS.docx")

doc = Document()

# ─── Page setup ───
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

# ─── Style definitions ───
style = doc.styles

# Name style
name_style = style.add_style('CVName', WD_STYLE_TYPE.PARAGRAPH)
name_style.font.name = 'Calibri'
name_style.font.size = Pt(16)
name_style.font.bold = True
name_style.font.color.rgb = RGBColor(0, 0, 0)
name_style.paragraph_format.space_before = Pt(0)
name_style.paragraph_format.space_after = Pt(2)
name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Contact style
contact_style = style.add_style('CVContact', WD_STYLE_TYPE.PARAGRAPH)
contact_style.font.name = 'Calibri'
contact_style.font.size = Pt(9)
contact_style.font.color.rgb = RGBColor(80, 80, 80)
contact_style.paragraph_format.space_before = Pt(0)
contact_style.paragraph_format.space_after = Pt(6)

# Section header style
section_style = style.add_style('CVSection', WD_STYLE_TYPE.PARAGRAPH)
section_style.font.name = 'Calibri'
section_style.font.size = Pt(11)
section_style.font.bold = True
section_style.font.color.rgb = RGBColor(0, 0, 0)
section_style.paragraph_format.space_before = Pt(8)
section_style.paragraph_format.space_after = Pt(3)
# Add bottom border
pPr = section_style.element.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/></w:pBdr>')
pPr.append(pBdr)

# Company/Role header
role_style = style.add_style('CVRole', WD_STYLE_TYPE.PARAGRAPH)
role_style.font.name = 'Calibri'
role_style.font.size = Pt(10)
role_style.font.bold = True
role_style.font.color.rgb = RGBColor(0, 0, 0)
role_style.paragraph_format.space_before = Pt(5)
role_style.paragraph_format.space_after = Pt(1)

# Sub-role (title + dates)
subrole_style = style.add_style('CVSubRole', WD_STYLE_TYPE.PARAGRAPH)
subrole_style.font.name = 'Calibri'
subrole_style.font.size = Pt(9.5)
subrole_style.font.color.rgb = RGBColor(60, 60, 60)
subrole_style.paragraph_format.space_before = Pt(0)
subrole_style.paragraph_format.space_after = Pt(2)

# Bullet style
bullet_style = style.add_style('CVBullet', WD_STYLE_TYPE.PARAGRAPH)
bullet_style.font.name = 'Calibri'
bullet_style.font.size = Pt(9)
bullet_style.font.color.rgb = RGBColor(30, 30, 30)
bullet_style.paragraph_format.space_before = Pt(1)
bullet_style.paragraph_format.space_after = Pt(1)
bullet_style.paragraph_format.left_indent = Inches(0.25)
bullet_style.paragraph_format.first_line_indent = Inches(-0.15)

# Skills style
skills_style = style.add_style('CVSkills', WD_STYLE_TYPE.PARAGRAPH)
skills_style.font.name = 'Calibri'
skills_style.font.size = Pt(9)
skills_style.font.color.rgb = RGBColor(30, 30, 30)
skills_style.paragraph_format.space_before = Pt(1)
skills_style.paragraph_format.space_after = Pt(1)

# ─── Helper functions ───
def add_section_header(text):
    doc.add_paragraph(text.upper(), style='CVSection')

def add_role(company, location):
    doc.add_paragraph(f"{company}  |  {location}", style='CVRole')

def add_subrole(title, dates):
    p = doc.add_paragraph(style='CVSubRole')
    run = p.add_run(title)
    run.bold = True
    p.add_run(f"  |  {dates}")

def add_bullet(text):
    p = doc.add_paragraph(style='CVBullet')
    p.add_run("\u2022  ")  # bullet character
    p.add_run(text)

def add_skills_line(category, items):
    p = doc.add_paragraph(style='CVSkills')
    run = p.add_run(f"{category}: ")
    run.bold = True
    p.add_run(items)

def add_education(school, location, degree, gpa_text, bullets=None):
    add_role(school, location)
    if gpa_text:
        add_subrole(degree, gpa_text)
    else:
        p = doc.add_paragraph(style='CVSubRole')
        run = p.add_run(degree)
        run.bold = True
    if bullets:
        for b in bullets:
            add_bullet(b)

# ═══════════════════════════════════════════════════
#                    CV CONTENT
# ═══════════════════════════════════════════════════

# ─── Name and contact ───
doc.add_paragraph("Tomas Maria Burnay Batalha", style='CVName')
p = doc.add_paragraph(style='CVContact')
p.add_run("Lisbon, Portugal (Open to Relocation to Dublin)  |  (+351) 936 124 118  |  tomas.b.batalha@gmail.com  |  linkedin.com/in/tomasmbatalha")

# ─── Professional Summary ───
add_section_header("Professional Summary")
p = doc.add_paragraph(style='CVSkills')
p.add_run(
    "Former Google intern and startup Co-Founder with full-cycle new business sales experience. "
    "Built a \u20ac147K B2B pipeline from zero through outbound prospecting, cold outreach, and enterprise deal closing. "
    "Managed a 2,311-account territory at Amazon, outperforming full-time colleagues. "
    "Seeking to bring builder mentality and commercial execution to Google's New Business Sales team."
)

# ─── Work Experience ───
add_section_header("Work Experience")

# Pairwire
add_role("Pairwire", "New York, USA (Remote)")
add_subrole("Co-Founder and Head of Sales", "February 2024 - December 2025")
add_bullet("Built \u20ac147K pipeline of contracted enterprise pilots from zero in 11 months pre-funding through outbound prospecting, cold calling, consulting, and relationship management")
add_bullet("Closed 23 enterprise engagements including Fortune 500 firms across 6-9 month full-cycle sales processes")
add_bullet("Designed an AI-powered outbound engine processing 7,000+ leads, automating 80% of prospecting via custom tech stack (Python, Expand.io, AI-enhanced workflows)")
add_bullet("Led end-to-end deal management from first touch to contract close, managing entire funnel with no sales team or brand recognition")
add_bullet("Created partnership and investor pitch materials; led negotiations with Google, Web Summit, and enterprise decision-makers")

# Google
add_role("Google", "Lisbon, Portugal")
add_subrole("Business Analyst Intern (Large Customer Sales)", "May 2024 - August 2024")
add_bullet("Led the \"Future Trends\" segment at Google's flagship partner event, presenting personally researched insights on Peak Season activation and media readiness to 237+ senior advertising clients")
add_bullet("Designed and automated three globally scalable dashboards using Apps Script and GMP tools to analyze ROAS vs. profit and CPA vs. conversions, projecting $1.41M in annual savings for media spend decisions")
add_bullet("Built a reporting platform for Partner Managers using Looker Studio, Analytics 360, and Apps Script, delivering an estimated \u20ac315K in annual efficiency gains")

# Amazon
add_role("Amazon", "Madrid, Spain")
add_subrole("Business Development Intern (Amazon Business)", "May 2022 - November 2022")
add_bullet("Led the B2B retail expansion into Portugal, managing a portfolio of 2,311 companies through prospecting, outreach, and account activation")
add_bullet("Outperformed KPI benchmarks by 21x in company account configurations during the Portugal B2B rollout")
add_bullet("Raised territory share from 56% below to 25% above team average, outperforming full-time colleagues in the region")
add_bullet("Compiled a 19-page strategic report on B2B eCommerce based on 8 company territories and 54 targeted client interviews")

# EY
add_role("EY", "Lisbon, Portugal")
add_subrole("Assurance Associate (Financial Services)", "February 2022 - April 2022")
add_bullet("Audited asset portfolios worth \u20ac41.3 billion across 3 insurance firms, developing rigorous data verification and attention to detail")

# ─── Education ───
add_section_header("Education")

add_education(
    "Rotterdam School of Management", "Rotterdam, Netherlands",
    "MSc in Strategic Management", "GPA: 8/10",
    ["Relevant Courses: Applied Machine Learning to Strategy (9.3/10), Python Fundamentals (8.3/10)"]
)

add_education(
    "National University of Singapore", "Singapore",
    "MSc Curriculum Term", "MBA-aligned coursework in Strategy and Innovation",
    ["Built a commercial and financial plan for a NUS PhD venture: 70% energy savings, 45% unit margins, 278% ROI in a $1.5B market"]
)

add_education(
    "Nova School of Business and Economics", "Lisbon, Portugal",
    "BSc in Management", "GPA: 16/20",
    ["Statistics (20/20), Procurement Management (19/20), Economics of the European Union (19/20)"]
)

# ─── Leadership & Activities ───
add_section_header("Leadership & Activities")
add_bullet("Nova Thirst Project - Founder & President: Led 23 volunteers, raised \u20ac2,539 for water access, organized 4 events including one with 35 NGOs (2020-2021)")
add_bullet("Nova Tech Club - Digital Transformation Consultant: Designed a job platform prototype and organized corporate events and workshops (2020-2021)")
add_bullet("Nova Social Consulting (WWF) - Pro Bono Consultant: Assessed organizational structure, identified 155 improvement areas, delivered 6 presentations to senior management (2020-2021)")

# ─── Skills ───
add_section_header("Skills")
add_skills_line("Sales and Marketing", "Google Ads, YouTube Ads, Google Display Network, CRM (Salesforce), LinkedIn Sales Navigator, Apollo, Full-Cycle Sales, Pipeline Management, Cold Outreach, Consultative Selling")
add_skills_line("Technical", "Python, Apps Script, Visual Basic for Applications (VBA), SQL, Looker Studio, Power BI, DAX")
add_skills_line("Certifications", "Google Ads Certified, Bloomberg Market Concepts, Microsoft Excel Specialist, TOEFL iBT")
add_skills_line("Languages", "Portuguese (Native), English (Fluent), Spanish (Fluent), French (Beginner)")

# ─── Save ───
doc.save(OUTPUT_PATH)
print(f"\u2705 Saved ATS-optimized CV to: {OUTPUT_PATH}")
print(f"   File size: {os.path.getsize(OUTPUT_PATH):,} bytes")
print()
print("This version has:")
print("  - ZERO tables (pure paragraphs + headings)")
print("  - Standard Word styles for ATS section detection")
print("  - Professional Summary included (wasn't in the table version)")
print("  - Clean single-column layout")
print("  - All text directly parseable by any ATS")
