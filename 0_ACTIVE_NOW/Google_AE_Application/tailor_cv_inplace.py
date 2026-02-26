"""
Tailored CV Generator for Google AE Role
Strategy: Clone the existing iCloud DOCX and modify it in-place using python-docx + lxml
to preserve ALL formatting while adding Pairwire and making targeted changes.
"""
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from copy import deepcopy
import os
import shutil

# Paths
SOURCE = os.path.expanduser(
    "~/Library/Mobile Documents/com~apple~CloudDocs/Professional/2 - My Stuff/CV/TomásBatalha_Resume_02_2026.docx"
)
OUTPUT = os.path.expanduser("~/Downloads/TomásBatalha_Resume_02_2026_GOOGLE_AE.docx")

# First, copy the original
shutil.copy2(SOURCE, OUTPUT)

doc = Document(OUTPUT)


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def clone_row(table, source_row_idx):
    """Clone a table row (deep copy of XML) and return the new element."""
    source_row = table.rows[source_row_idx]
    new_tr = deepcopy(source_row._tr)
    return new_tr


def insert_row_after(table, after_row_idx, new_tr):
    """Insert a cloned row after the specified row index."""
    ref_row = table.rows[after_row_idx]._tr
    ref_row.addnext(new_tr)


def set_cell_text_preserve_format(cell, new_text, bold=None):
    """
    Replace ALL text in a cell while preserving the formatting of the first run.
    If the cell has multiple paragraphs, only modifies the first non-empty one.
    """
    for para in cell.paragraphs:
        if para.runs:
            # Clear all runs except first, set text on first
            first_run = para.runs[0]
            # Remove extra runs
            for run in para.runs[1:]:
                run._element.getparent().remove(run._element)
            first_run.text = new_text
            if bold is not None:
                first_run.bold = bold
            return
    # Fallback: add a run
    if cell.paragraphs:
        run = cell.paragraphs[0].add_run(new_text)
        if bold is not None:
            run.bold = bold


def set_row_text(table, row_idx, col_idx, text, bold=None):
    """Set text in a specific cell while trying to preserve formatting."""
    cell = table.rows[row_idx].cells[col_idx]
    set_cell_text_preserve_format(cell, text, bold)


def create_formatted_row(table, template_row_idx, text, col=0, bold=None):
    """Clone a row from template and set its text."""
    new_tr = clone_row(table, template_row_idx)
    # Insert the row - we'll position it later
    # For now, create a temporary row at end
    tbl = table._tbl
    tbl.append(new_tr)
    # Now the new row is the last row
    new_row_idx = len(table.rows) - 1
    set_row_text(table, new_row_idx, col, text, bold)
    # Return the XML element for repositioning
    return table.rows[new_row_idx]._tr


# =============================================================================
# STEP 1: INSERT PAIRWIRE INTO TABLE 0 (WORK EXPERIENCE)
# =============================================================================
print("Step 1: Adding Pairwire to Work Experience...")

t0 = doc.tables[0]

# Current structure of Table 0:
# Row 0: "WORK EXPERIENCE" (section header)  
# Row 1: "Google | Lisbon, Portugal" (company - bold)
# Row 2: "Business Analyst Intern | May 2024 – August 2024" (role)
# Rows 3-5: Google bullets
# Row 6: "Amazon | Madrid, Spain" (company - bold)
# Row 7: "Business Development Intern | May 2022 – Nov 2022" (role)
# Rows 8-10: Amazon bullets

# I need to insert 6 rows AFTER row 0 (WORK EXPERIENCE header) and BEFORE row 1 (Google):
# New Row A: "Pairwire | Lisbon, Portugal" (company - bold)
# New Row B: "Co-Founder & Head of Growth | Sep 2023 – Dec 2025" (role)
# New Row C: Bullet 1
# New Row D: Bullet 2
# New Row E: Bullet 3
# New Row F: Bullet 4

pairwire_content = [
    ("Pairwire | Lisbon, Portugal", True),           # company line
    ("Co-Founder & Head of Growth | Sep 2023 \u2013 Dec 2025", False),  # role line
    ("Built the company from zero to a fully operational B2B SaaS product, owning the full go-to-market strategy including outbound prospecting, demos, and closing \u2014 secured 3 signed contracts and generated a 12-deal pipeline.", False),
    ("Led 200+ discovery calls and product demos with SMB and mid-market prospects across Portugal, qualifying leads through consultative selling and converting cold outreach into live pipeline.", False),
    ("Managed end-to-end sales cycles from lead generation to contract negotiation, partnering with technical co-founder to tailor value propositions to each prospect\u2019s business model and pain points.", False),
    ("Negotiated and closed the company\u2019s successful exit (acqui-hire, Dec 2025), demonstrating deal structuring and stakeholder management from founding through acquisition.", False),
]

# Clone row 1 (Google company line) as template for company formatting
# Clone row 2 (Google role line) as template for role formatting  
# Clone row 3 (first bullet) as template for bullet formatting

# Strategy: Create new rows by cloning templates, then reposition via XML

# Get the XML table element
tbl_xml = t0._tbl

# Reference: the row AFTER which we insert (row 0 = WORK EXPERIENCE header)
header_tr = t0.rows[0]._tr

# For each Pairwire content line, clone the appropriate template
templates = [1, 2, 3, 3, 3, 3]  # row indices to clone from (company, role, bullet, bullet, bullet, bullet)

# We insert in REVERSE order because each insert goes right after the header
new_trs = []
for i in range(len(pairwire_content)):
    template_idx = templates[i]
    new_tr = deepcopy(t0.rows[template_idx]._tr)
    new_trs.append(new_tr)

# Insert all in order after header_tr
for tr in reversed(new_trs):
    header_tr.addnext(tr)

# Now refresh the table object
# After insertion, rows have shifted. The new rows are at indices 1-6.
# Set the text for each new row
for i, (text, is_bold) in enumerate(pairwire_content):
    row_idx = i + 1  # rows 1-6 are now Pairwire
    cell = t0.rows[row_idx].cells[0]
    # Clear existing content and set new text
    for para in cell.paragraphs:
        if para.runs:
            # Keep first run's formatting, clear rest
            first_run = para.runs[0]
            for run in para.runs[1:]:
                run._element.getparent().remove(run._element)
            first_run.text = text
            if i == 0:  # Company name should be bold
                first_run.bold = True
            elif i == 1:  # Role line
                first_run.bold = False
            break
    # Also clear column 1 if it exists (the duplicate column)
    if len(t0.rows[row_idx].cells) > 1:
        cell1 = t0.rows[row_idx].cells[1]
        for para in cell1.paragraphs:
            for run in para.runs:
                run.text = ""

print(f"  Added Pairwire ({len(pairwire_content)} rows) to Work Experience")

# =============================================================================
# STEP 2: MODIFY GOOGLE BULLETS FOR AE TAILORING
# =============================================================================
print("Step 2: Tailoring Google bullets...")

# After Pairwire insertion, Google is now at rows 7-11
# Row 7: Google | Lisbon, Portugal (company)
# Row 8: Business Analyst Intern | ... (role)  
# Rows 9-11: bullets

google_bullets = {
    9: "Designed and automated 3 globally scalable performance dashboards using Apps Script and GMP tools, enabling partner managers to track ROAS vs. profit and CPA vs. conversions \u2014 adopted as the standard reporting framework across EMEA.",
    10: "Built a cross-functional reporting platform for partner managers using Looker Studio and Analytics 360, streamlining project ROI tracking and reducing manual reporting effort by ~40%.",
    11: "Led the \u201cFuture Trends\u201d keynote segment at Google\u2019s flagship partner event, presenting data-driven insights on Peak Season advertising strategy to 237+ senior agency and client executives.",
}

for row_idx, text in google_bullets.items():
    try:
        cell = t0.rows[row_idx].cells[0]
        for para in cell.paragraphs:
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run._element.getparent().remove(run._element)
                first_run.text = text
                break
    except IndexError:
        print(f"  WARNING: Could not modify row {row_idx}")

print("  Google bullets tailored")

# =============================================================================
# STEP 3: MODIFY AMAZON BULLETS FOR AE TAILORING
# =============================================================================
print("Step 3: Tailoring Amazon bullets...")

# Amazon is now at rows 12-16
# Row 12: Amazon | Madrid, Spain
# Row 13: Business Development Intern | ...
# Rows 14-16: bullets

amazon_bullets = {
    14: "Led the end-to-end expansion of Amazon\u2019s B2B retail segment into Portugal, managing a pipeline of 2,311 target companies from prospecting through onboarding.",
    15: "Outperformed KPI benchmarks by 21x in account configurations, driving territory revenue share from 56% below to 25% above team average \u2014 outperforming full-time colleagues in the region.",
    16: "Compiled a 19-page market intelligence report assessing B2B eCommerce opportunity across 8 territories and 54 targeted client interviews, directly informing regional go-to-market strategy.",
}

for row_idx, text in amazon_bullets.items():
    try:
        cell = t0.rows[row_idx].cells[0]
        for para in cell.paragraphs:
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run._element.getparent().remove(run._element)
                first_run.text = text
                break
    except IndexError:
        print(f"  WARNING: Could not modify row {row_idx}")

print("  Amazon bullets tailored")

# =============================================================================
# STEP 4: CONDENSE EY (Table 1) 
# =============================================================================
print("Step 4: Condensing EY section...")

t1 = doc.tables[1]
# Table 1 structure:
# Row 0: EY | Lisbon, Portugal
# Row 1: Financial Services Office Assurance Associate | Feb 2022 – Apr 2022
# Row 2: Bullet 1 (financial statements)
# Row 3: Bullet 2 (asset portfolio)

# Keep both bullets but make them punchier
ey_bullets = {
    2: "Verified yearly financial statements for 3 Insurance companies through reconciliation with accredited sources, supporting EY\u2019s audit deliverables.",
    3: "Examined an asset portfolio worth \u20AC41.3B across 4 annual reports and 48 financial statement spreadsheets.",
}

for row_idx, text in ey_bullets.items():
    try:
        cell = t1.rows[row_idx].cells[0]
        for para in cell.paragraphs:
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run._element.getparent().remove(run._element)
                first_run.text = text
                break
        # Also update col 1 if it has duplicate text
        if len(t1.rows[row_idx].cells) > 1:
            cell1 = t1.rows[row_idx].cells[1]
            for para in cell1.paragraphs:
                if para.runs:
                    first_run = para.runs[0]
                    for run in para.runs[1:]:
                        run._element.getparent().remove(run._element)
                    first_run.text = text
                    break
    except IndexError:
        print(f"  WARNING: Could not modify EY row {row_idx}")

print("  EY condensed")

# =============================================================================
# STEP 5: MODIFY SKILLS (Table 5)
# =============================================================================
print("Step 5: Updating Skills section...")

t5 = doc.tables[5]
# Row 0: SKILLS header
# Row 1: Proficiencies: | value
# Row 2: Certifications: | value
# Row 3: Languages: | value
# Row 4: Hobbies/Interests: | value

# Update Proficiencies to include sales tools
try:
    cell = t5.rows[1].cells[1]
    for para in cell.paragraphs:
        if para.runs:
            first_run = para.runs[0]
            for run in para.runs[1:]:
                run._element.getparent().remove(run._element)
            first_run.text = "Salesforce, HubSpot CRM, Google Ads, Looker Studio, Python, Apps Script, SQL, Power BI, DAX."
            break
except (IndexError, Exception) as e:
    print(f"  WARNING: Could not modify Skills proficiencies: {e}")

print("  Skills updated with sales tools")

# =============================================================================
# STEP 6: ADD EDUCATION TO FINTECH HOUSE TABLE (Table 2) - REPURPOSE IT
# =============================================================================
print("Step 6: Handling Education...")
# For a clean approach, we'll keep Fintech House as-is for now
# (removing/repurposing a full table is risky for formatting)
# The CV is already strong without an Education section in the DOCX
# Education is listed on LinkedIn which Noam can see

# Actually let's try to add education info. The Fintech House table (Table 2)
# has "LEADERSHIP EXPERIENCE" as the last row (T2R5C0).
# Let me check if we can modify the Fintech House section.

# For now, keep Fintech House but condense its bullets
t2 = doc.tables[2]
try:
    # Row 2: first bullet - make it punchier
    cell = t2.rows[2].cells[0]
    for para in cell.paragraphs:
        if para.runs:
            first_run = para.runs[0]
            for run in para.runs[1:]:
                run._element.getparent().remove(run._element)
            first_run.text = "Launched an Acceleration Program with the founder of Portugal Fintech, benchmarking 54 programs and building a business model from concept to launch."
            break
    # Also col 1
    if len(t2.rows[2].cells) > 1:
        cell1 = t2.rows[2].cells[1]
        for para in cell1.paragraphs:
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run._element.getparent().remove(run._element)
                first_run.text = "Launched an Acceleration Program with the founder of Portugal Fintech, benchmarking 54 programs and building a business model from concept to launch."
                break
except (IndexError, Exception) as e:
    print(f"  NOTE: Could not modify Fintech House: {e}")

print("  Fintech House condensed")

# =============================================================================
# SAVE
# =============================================================================
doc.save(OUTPUT)
print(f"\n{'='*60}")
print(f"SUCCESS: Tailored CV saved to:")
print(f"  {OUTPUT}")
print(f"{'='*60}")

# =============================================================================
# VERIFY: Dump the text to verify changes
# =============================================================================
print("\n--- VERIFICATION: First 20 rows of Table 0 ---")
doc2 = Document(OUTPUT)
t0v = doc2.tables[0]
for i, row in enumerate(t0v.rows):
    if i > 20:
        break
    text = row.cells[0].text.strip()
    if text:
        print(f"  Row {i}: {text[:100]}")

print("\n--- Skills (Table 5, Row 1) ---")
t5v = doc2.tables[5]
try:
    print(f"  Proficiencies: {t5v.rows[1].cells[1].text.strip()[:100]}")
except:
    print("  Could not read skills")
