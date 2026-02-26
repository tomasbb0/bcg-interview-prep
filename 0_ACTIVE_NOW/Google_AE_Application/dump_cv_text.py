"""Dump ALL text content from the iCloud CV to see what's there"""
from docx import Document
import os

cv_path = os.path.expanduser(
    "~/Library/Mobile Documents/com~apple~CloudDocs/Professional/2 - My Stuff/CV/TomásBatalha_Resume_02_2026.docx"
)

doc = Document(cv_path)

print("=" * 80)
print("ALL PARAGRAPHS (outside tables)")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else "None"
        bold_runs = [r.text for r in para.runs if r.bold]
        print(f"[P{i}] ({style}) {para.text[:120]}")

print("\n" + "=" * 80)
print("ALL TABLES")
print("=" * 80)
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {t_idx} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                # Truncate long text but show structure
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        print(f"  T{t_idx}R{r_idx}C{c_idx}: {line[:150]}")
