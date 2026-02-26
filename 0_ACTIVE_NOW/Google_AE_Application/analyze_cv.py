#!/usr/bin/env python3
"""Analyze the iCloud CV structure for formatting preservation."""

from docx import Document

ICLOUD_CV = '/Users/tomasbatalha/Library/Mobile Documents/com~apple~CloudDocs/Professional/2 - My Stuff/CV/TomásBatalha_Resume_02_2026.docx'
doc = Document(ICLOUD_CV)

print('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[P{i}] style={p.style.name}: {p.text[:80]}')
        for ri, r in enumerate(p.runs):
            if r.text.strip():
                print(f'  run{ri}: text={r.text[:40]!r} bold={r.bold} italic={r.italic} size={r.font.size} font={r.font.name}')

print('\n=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'\nTable {ti} ({len(table.rows)} rows, {len(table.columns)} cols):')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for pi, p in enumerate(cell.paragraphs):
                if p.text.strip():
                    print(f'  [{ti}][R{ri}][C{ci}][P{pi}] style={p.style.name}: {p.text[:70]}')
                    for rri, r in enumerate(p.runs):
                        if r.text.strip():
                            print(f'    run{rri}: {r.text[:50]!r} bold={r.bold} size={r.font.size}')
