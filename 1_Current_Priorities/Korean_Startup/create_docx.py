from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set margins (1 inch = 2.54 cm)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Helper functions
def add_heading_centered(text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    if level == 0:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(14)
    return p

def add_heading_left(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.space_after = Pt(12)
    return p

def add_para(text, bold_start=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_start:
        run = p.add_run(bold_start)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        rest = text[len(bold_start):]
        run2 = p.add_run(rest)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        row.cells[i].text = cell_text
        for para in row.cells[i].paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

def add_line():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 80)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ============ DOCUMENT CONTENT ============

# Title
add_heading_centered('INDEPENDENT CONTRACTOR SERVICES AGREEMENT')
doc.add_paragraph()

# Header info
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Agreement Reference: SA-2024-001')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Effective Date: _______________________')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

add_line()
doc.add_paragraph()

# PARTIES
add_heading_left('PARTIES')

add_para('THE COMPANY', 'THE COMPANY')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Company Name', '________________________________________________'])
add_table_row(table, ['SIRET / Registration No.', '________________________________________________'])
add_table_row(table, ['Registered Address', '________________________________________________'])
add_table_row(table, ['Country of Incorporation', '________________________________________________'])
add_table_row(table, ['Represented By', '________________________________________________'])
add_table_row(table, ['Position', '________________________________________________'])
add_table_row(table, ['Email', '________________________________________________'])
doc.add_paragraph()

add_para('Hereinafter referred to as "the Company" or "Client"')
doc.add_paragraph()

add_para('THE CONTRACTOR', 'THE CONTRACTOR')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Full Name', 'Tomás Maria Burnay Batalha'])
add_table_row(table, ['Portuguese NIF', '________________________________________________'])
add_table_row(table, ['Address', '________________________________________________'])
add_table_row(table, ['Country of Tax Residence', 'Portugal'])
add_table_row(table, ['Email', '________________________________________________'])
add_table_row(table, ['Phone', '________________________________________________'])
doc.add_paragraph()

add_para('Hereinafter referred to as "the Contractor" or "Service Provider"')

add_line()
doc.add_paragraph()

# RECITALS
add_heading_left('RECITALS')

add_para('WHEREAS:', 'WHEREAS:')
doc.add_paragraph()

add_para('A. The Company operates a healthtech business and seeks to expand its commercial operations through improved go-to-market capabilities;')
add_para('B. The Contractor is an independent professional with expertise in commercial strategy, software development, and growth operations;')
add_para('C. The Contractor operates as a self-employed professional (trabalhador independente) under Portuguese law and issues invoices via the Portuguese Tax Authority\'s Recibos Verdes system;')
add_para('D. Both parties wish to establish a clear framework for the provision of services, compensation, and mutual obligations;')
doc.add_paragraph()

add_para('NOW, THEREFORE, the parties agree as follows:', 'NOW, THEREFORE,')

add_line()
doc.add_paragraph()

# ARTICLE 1
add_heading_left('ARTICLE 1: NATURE OF RELATIONSHIP')

add_para('1.1 Independent Contractor Status. The Contractor is engaged as an independent contractor (prestataire de services / prestador de serviços) and not as an employee. There is no relationship of subordination (lien de subordination) between the parties.', '1.1 Independent Contractor Status.')
doc.add_paragraph()

add_para('1.2 No Employment Relationship. This Agreement does not create an employment contract (contrat de travail / contrato de trabalho). The Contractor:', '1.2 No Employment Relationship.')
add_bullet('Determines their own working hours and methods')
add_bullet('Provides their own equipment and workspace')
add_bullet('Is responsible for their own taxes, social security, and insurance')
add_bullet('May provide services to other clients')
doc.add_paragraph()

add_para('1.3 EU Services Freedom. This Agreement is entered into under the freedom to provide services within the European Union, in accordance with Articles 56-62 of the Treaty on the Functioning of the European Union (TFEU).', '1.3 EU Services Freedom.')

doc.add_paragraph()

# ARTICLE 2
add_heading_left('ARTICLE 2: SCOPE OF SERVICES')

add_para('2.1 Primary Services. The Contractor agrees to provide the following services:', '2.1 Primary Services.')
doc.add_paragraph()

add_para('A) GTM System Development', 'A) GTM System Development')
add_bullet('Custom Go-To-Market web application')
add_bullet('AI-personalized bulk email campaign functionality')
add_bullet('Lead import and management system')
add_bullet('Sales pipeline visualization dashboard')
add_bullet('Email open/click tracking and analytics')
add_bullet('Campaign scheduling and management interface')
doc.add_paragraph()

add_para('B) Commercial Growth Services', 'B) Commercial Growth Services')
add_bullet('Outbound lead generation and prospecting')
add_bullet('Sales campaign design and execution')
add_bullet('Pipeline management and optimization')
add_bullet('Performance reporting')
doc.add_paragraph()

add_para('2.2 Minimum Viable Product (MVP). The GTM system MVP shall include the core features listed in 2.1(A) in functional, deployable form suitable for business operations.', '2.2 Minimum Viable Product (MVP).')
doc.add_paragraph()

add_para('2.3 Time Commitment. The Contractor commits to approximately fifteen to twenty (15-20) hours per week during the term of this Agreement.', '2.3 Time Commitment.')
doc.add_paragraph()

add_para('2.4 Exclusions. Unless separately agreed in writing, Services do not include: hardware procurement, third-party software licenses, legal or accounting advice, or in-person attendance at Company premises.', '2.4 Exclusions.')

doc.add_paragraph()

# ARTICLE 3
add_heading_left('ARTICLE 3: COMPENSATION')

add_para('3.1 Advance Payment. Upon signing this Agreement, the Company shall pay:', '3.1 Advance Payment.')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Signing Advance', '€ _____________'])
add_table_row(table, ['Tool & Infrastructure Budget', '€ 150.00'])
add_table_row(table, ['Total Upfront Payment', '€ _____________'])
doc.add_paragraph()

add_para('3.2 Work Equivalent. The Signing Advance represents pre-payment for approximately sixty (60) hours of work at an effective rate of €_____________ per hour.', '3.2 Work Equivalent.')
doc.add_paragraph()

add_para('3.3 Monthly Compensation. Following the advance period:', '3.3 Monthly Compensation.')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Monthly Retainer', '€ _____________ /month'])
add_table_row(table, ['Commission on Closed Deals', '_______% of net deal value'])
doc.add_paragraph()

add_para('3.4 Payment Terms.', '3.4 Payment Terms.')
add_bullet('Advance: Due within three (3) business days of signing')
add_bullet('Monthly Retainer: Due on the 1st of each calendar month')
add_bullet('Commission: Due within fifteen (15) days of Company receiving deal payment')
doc.add_paragraph()

add_para('3.5 Running Costs. The Company shall reimburse ongoing operational costs for the GTM system, estimated at €50/month, covering hosting, email infrastructure, and AI API usage.', '3.5 Running Costs.')
doc.add_paragraph()

add_para('3.6 Payment Method. All payments shall be made via bank transfer to the Contractor\'s Portuguese bank account:', '3.6 Payment Method.')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Account Holder', 'Tomás Maria Burnay Batalha'])
add_table_row(table, ['IBAN', '________________________________________________'])
add_table_row(table, ['BIC/SWIFT', '________________________________________________'])
add_table_row(table, ['Bank Name', '________________________________________________'])

doc.add_paragraph()

# ARTICLE 4
add_heading_left('ARTICLE 4: INVOICING AND TAX COMPLIANCE')

add_para('4.1 Invoicing. The Contractor shall issue invoices (Fatura-Recibo) via the Portuguese Tax Authority portal (Portal das Finanças) for all amounts due.', '4.1 Invoicing.')
doc.add_paragraph()

add_para('4.2 Invoice Requirements. Each invoice shall include:', '4.2 Invoice Requirements.')
add_bullet('Contractor\'s name, address, and Portuguese NIF')
add_bullet('Company\'s name, address, and SIRET/registration number')
add_bullet('Description of services rendered')
add_bullet('Amount in EUR')
add_bullet('VAT notation as per Article 4.3')
doc.add_paragraph()

add_para('4.3 VAT Treatment. As an intra-EU B2B service:', '4.3 VAT Treatment.')
add_bullet('If Contractor is VAT-exempt: Invoice states "IVA – regime de isenção, artigo 53.º do CIVA"')
add_bullet('If reverse charge applies: Invoice states "IVA – autoliquidação / Reverse charge – Article 196, EU VAT Directive 2006/112/EC"')
doc.add_paragraph()

add_para('4.4 Tax Responsibility. The Contractor is solely responsible for:', '4.4 Tax Responsibility.')
add_bullet('Portuguese income tax (IRS) obligations')
add_bullet('Portuguese social security (Segurança Social) contributions')
add_bullet('Any applicable VAT obligations')
doc.add_paragraph()

add_para('4.5 No Withholding. The Company shall pay gross amounts without withholding, as no withholding tax applies to EU cross-border B2B services between Portugal and France.', '4.5 No Withholding.')

doc.add_paragraph()

# ARTICLE 5
add_heading_left('ARTICLE 5: DELIVERABLES AND TIMELINE')

add_para('5.1 MVP Delivery Deadline. The Contractor shall deliver a functional GTM System MVP within six (6) weeks of receiving the Advance Payment.', '5.1 MVP Delivery Deadline.')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Advance Payment Received', '_______________________ (Date)'])
add_table_row(table, ['MVP Deadline', '_______________________ (Date)'])
doc.add_paragraph()

add_para('5.2 Progress Reporting. The Contractor shall provide weekly written updates including:', '5.2 Progress Reporting.')
add_bullet('Hours worked during the period')
add_bullet('Tasks completed')
add_bullet('Tasks planned for the following week')
add_bullet('Blockers, risks, or concerns')
doc.add_paragraph()

add_para('5.3 Acceptance Process. The Company shall have seven (7) calendar days following MVP delivery to:', '5.3 Acceptance Process.')
add_bullet('Review the deliverables')
add_bullet('Request reasonable modifications')
add_bullet('Provide written acceptance or objections')
doc.add_paragraph()

add_para('If no written response is received within seven (7) days, the MVP shall be deemed accepted.')

doc.add_paragraph()

# ARTICLE 6
add_heading_left('ARTICLE 6: ADVANCE REFUND CONDITIONS')

add_para('6.1 Refund Triggers (Contractor Responsibility). The Contractor shall refund a proportionate amount of the Advance if:', '6.1 Refund Triggers (Contractor Responsibility).')
add_bullet('(a) Contractor ceases communication for fourteen (14) or more consecutive days without valid reason;')
add_bullet('(b) Contractor explicitly terminates before MVP delivery without mutual consent;')
add_bullet('(c) Contractor fails to deliver MVP by deadline (excluding Company-caused delays).')
doc.add_paragraph()

add_para('6.2 Refund Calculation.', '6.2 Refund Calculation.')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Refund Amount = Advance Payment − (Hours Worked × Hourly Rate)')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.italic = True
doc.add_paragraph()

add_para('Hours Worked shall be determined by the weekly progress reports under Article 5.2.')
doc.add_paragraph()

add_para('6.3 No Refund Conditions. The Contractor retains the full Advance with no refund obligation if:', '6.3 No Refund Conditions.')
add_bullet('(a) MVP is delivered as specified;')
add_bullet('(b) Company terminates for reasons other than Contractor\'s material breach;')
add_bullet('(c) Company fails to respond for fourteen (14) or more consecutive days;')
add_bullet('(d) Company materially changes scope without written agreement;')
add_bullet('(e) Mutual agreement to end after equivalent hours worked;')
add_bullet('(f) Force majeure, illness, or emergency (with prompt communication).')
doc.add_paragraph()

add_para('6.4 Refund Timeline. Any refund due shall be paid within thirty (30) calendar days of the triggering event.', '6.4 Refund Timeline.')

doc.add_paragraph()

# ARTICLE 7
add_heading_left('ARTICLE 7: INTELLECTUAL PROPERTY')

add_para('7.1 Assignment. Upon full payment of all compensation due under this Agreement, all intellectual property rights in the GTM system and deliverables created under this Agreement shall transfer to the Company.', '7.1 Assignment.')
doc.add_paragraph()

add_para('7.2 Pre-Existing IP. The Contractor retains all rights to pre-existing tools, frameworks, libraries, and methodologies. The Company receives a perpetual, non-exclusive license to use any pre-existing IP incorporated in the deliverables.', '7.2 Pre-Existing IP.')
doc.add_paragraph()

add_para('7.3 Interim License. Until full payment is received, the Company receives a limited, revocable license to use the deliverables for internal business purposes only.', '7.3 Interim License.')

doc.add_paragraph()

# ARTICLE 8
add_heading_left('ARTICLE 8: CONFIDENTIALITY')

add_para('8.1 Confidential Information. Both parties agree to maintain the confidentiality of proprietary information disclosed during the engagement, including: business strategies, customer data, technical specifications, pricing, and financial information.', '8.1 Confidential Information.')
doc.add_paragraph()

add_para('8.2 Duration. Confidentiality obligations survive termination for two (2) years.', '8.2 Duration.')
doc.add_paragraph()

add_para('8.3 Permitted Disclosures. Confidentiality obligations do not apply to information that:', '8.3 Permitted Disclosures.')
add_bullet('(a) Becomes publicly available through no fault of the receiving party;')
add_bullet('(b) Was already known to the receiving party prior to disclosure;')
add_bullet('(c) Is independently developed by the receiving party;')
add_bullet('(d) Is required to be disclosed by law or court order.')

doc.add_paragraph()

# ARTICLE 9
add_heading_left('ARTICLE 9: GOOD FAITH AND DISPUTE RESOLUTION')

add_para('9.1 Good Faith. Both parties commit to acting in good faith and communicating openly about any concerns or issues.', '9.1 Good Faith.')
doc.add_paragraph()

add_para('9.2 Notice Before Breach Claim. Neither party may claim material breach without first providing seven (7) days written notice specifying:', '9.2 Notice Before Breach Claim.')
add_bullet('The alleged breach')
add_bullet('The remedy requested')
add_bullet('A reasonable opportunity to cure')
doc.add_paragraph()

add_para('9.3 Escalation Process.', '9.3 Escalation Process.')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
add_table_row(table, ['Step 1: Direct Discussion', '14 days – Parties attempt to resolve directly'])
add_table_row(table, ['Step 2: Mediation', '30 days – Neutral third-party mediation'])
add_table_row(table, ['Step 3: Legal Action', 'If mediation fails – Court proceedings'])
doc.add_paragraph()

add_para('9.4 Jurisdiction. Any legal proceedings shall be brought before the courts of:', '9.4 Jurisdiction.')
add_bullet('☐ Paris, France')
add_bullet('☐ Lisbon, Portugal')
add_bullet('☐ Other: _________________________________')
doc.add_paragraph()
add_para('(Tick one by mutual agreement)')

doc.add_paragraph()

# ARTICLE 10
add_heading_left('ARTICLE 10: TERM AND TERMINATION')

add_para('10.1 Initial Term. This Agreement begins on the Effective Date and continues for three (3) months.', '10.1 Initial Term.')
doc.add_paragraph()

add_para('10.2 Renewal. After the initial term, this Agreement renews month-to-month unless either party provides fourteen (14) days written notice of non-renewal.', '10.2 Renewal.')
doc.add_paragraph()

add_para('10.3 Termination for Convenience. Either party may terminate with fourteen (14) days written notice, subject to the refund provisions in Article 6.', '10.3 Termination for Convenience.')
doc.add_paragraph()

add_para('10.4 Termination for Cause. Either party may terminate immediately upon material breach, following the notice procedure in Article 9.2.', '10.4 Termination for Cause.')
doc.add_paragraph()

add_para('10.5 Effect of Termination. Upon termination:', '10.5 Effect of Termination.')
add_bullet('All unpaid amounts become immediately due')
add_bullet('Confidentiality obligations continue per Article 8.2')
add_bullet('IP provisions apply per Article 7')

doc.add_paragraph()

# ARTICLE 11
add_heading_left('ARTICLE 11: GENERAL PROVISIONS')

add_para('11.1 Language. This Agreement is drafted in English by mutual consent of the parties.', '11.1 Language.')
doc.add_paragraph()

add_para('11.2 Governing Law. This Agreement shall be governed by and construed in accordance with:', '11.2 Governing Law.')
add_bullet('☐ French law (droit français)')
add_bullet('☐ Portuguese law (lei portuguesa)')
add_bullet('☐ Other: _________________________________')
doc.add_paragraph()
add_para('(Tick one by mutual agreement)')
doc.add_paragraph()

add_para('11.3 Entire Agreement. This Agreement constitutes the entire agreement between the parties and supersedes all prior negotiations, representations, and agreements.', '11.3 Entire Agreement.')
doc.add_paragraph()

add_para('11.4 Amendments. Any amendments must be in writing and signed by both parties.', '11.4 Amendments.')
doc.add_paragraph()

add_para('11.5 Severability. If any provision is found unenforceable, the remaining provisions continue in full effect.', '11.5 Severability.')
doc.add_paragraph()

add_para('11.6 Notices. All notices shall be in writing and sent via email to the addresses specified in the Parties section, with read receipt requested.', '11.6 Notices.')
doc.add_paragraph()

add_para('11.7 Counterparts. This Agreement may be executed in counterparts, including electronic signatures, each of which shall constitute an original.', '11.7 Counterparts.')

add_line()
doc.add_paragraph()

# SIGNATURES
add_heading_left('SIGNATURES')

add_para('By signing below, both parties confirm they have read, understood, and agree to be bound by the terms of this Agreement.')
doc.add_paragraph()
doc.add_paragraph()

add_para('THE COMPANY', 'THE COMPANY')
doc.add_paragraph()

add_para('Signature:          ________________________________________________')
add_para('Printed Name:       ________________________________________________')
add_para('Title:              ________________________________________________')
add_para('Date:               ________________________________________________')
add_para('Place:              ________________________________________________')

doc.add_paragraph()
doc.add_paragraph()

add_para('THE CONTRACTOR', 'THE CONTRACTOR')
doc.add_paragraph()

add_para('Signature:          ________________________________________________')
add_para('Printed Name:       Tomás Maria Burnay Batalha')
add_para('Date:               ________________________________________________')
add_para('Place:              Portugal')

add_line()
doc.add_paragraph()

# EXHIBIT A
add_heading_centered('EXHIBIT A: PAYMENT SCHEDULE', level=1)
doc.add_paragraph()

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Payment'
hdr[1].text = 'Amount (EUR)'
hdr[2].text = 'Due Date'
hdr[3].text = 'Status'
add_table_row(table, ['Signing Advance', '€ _________', 'Upon signing', '☐ Paid'])
add_table_row(table, ['Tool Budget', '€ 150.00', 'Upon signing', '☐ Paid'])
add_table_row(table, ['Month 1 Retainer', '€ _________', '_____________', '☐ Paid'])
add_table_row(table, ['Month 2 Retainer', '€ _________', '_____________', '☐ Paid'])
add_table_row(table, ['Month 3 Retainer', '€ _________', '_____________', '☐ Paid'])
add_table_row(table, ['Running Costs', '~€ 50.00', '1st of month', '☐ Paid'])

doc.add_paragraph()

# EXHIBIT B
add_heading_centered('EXHIBIT B: GTM SYSTEM MVP SPECIFICATIONS', level=1)
doc.add_paragraph()

add_para('The Minimum Viable Product shall include:')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Description'
hdr[2].text = 'Priority'
add_table_row(table, ['Lead Management', 'Import CSV/Excel, search, filter, tag leads', 'Must Have'])
add_table_row(table, ['Email Campaigns', 'Create, schedule, send bulk personalized emails', 'Must Have'])
add_table_row(table, ['AI Personalization', 'GPT-powered email customization per lead', 'Must Have'])
add_table_row(table, ['Pipeline Dashboard', 'Visual Kanban board for sales stages', 'Must Have'])
add_table_row(table, ['Email Tracking', 'Open and click tracking with analytics', 'Must Have'])
add_table_row(table, ['User Authentication', 'Secure login for Company team members', 'Must Have'])
add_table_row(table, ['Basic Reporting', 'Campaign performance metrics', 'Should Have'])

add_line()
doc.add_paragraph()

# End
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— END OF AGREEMENT —')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document Reference: SA-2024-001 | Prepared December 2024')
run.font.name = 'Times New Roman'
run.font.size = Pt(9)
run.italic = True

# Save
doc.save('/Users/tomasbatalha/Downloads/Tomas_Batalha_Future_Plan/1_Current_Priorities/Korean_Startup/Services_Agreement_FINAL.docx')
print('✅ Created Services_Agreement_FINAL.docx')
