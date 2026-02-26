"""
Generate Services Agreement in standard legal format:
- No tables (use indented text blocks)
- No bullet points with balls (use (a), (b), (c) numbering)
- Proper legal hierarchical numbering: Article → 1.1 → (a)
- Times New Roman, 11pt body, 14pt headings, 1 inch margins
"""

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set margins (1 inch all around)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

# ============ HELPER FUNCTIONS ============

def set_font(run, size=11, bold=False, italic=False):
    """Apply Times New Roman with specified formatting."""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_title(text):
    """Add centered title in bold 14pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_centered(text, size=11, bold=False, italic=False):
    """Add centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size, bold, italic)
    return p

def add_article(text):
    """Add article heading in bold 12pt with space before."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p

def add_section(number, title, content):
    """Add numbered section: 1.1 Title. Content text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    # Bold number and title
    run1 = p.add_run(f"{number} {title}. ")
    set_font(run1, bold=True)
    
    # Normal content
    run2 = p.add_run(content)
    set_font(run2)
    return p

def add_para(text, indent=0, first_indent=0.5):
    """Add justified paragraph with optional indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    run = p.add_run(text)
    set_font(run)
    return p

def add_plain(text, bold_prefix=None):
    """Add plain paragraph, optionally with bold prefix."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        set_font(run1, bold=True)
        run2 = p.add_run(text[len(bold_prefix):])
        set_font(run2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def add_subitem(letter, text, indent=0.75):
    """Add lettered sub-item: (a) text"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(f"({letter})  {text}")
    set_font(run)
    return p

def add_roman_item(numeral, text, indent=0.75):
    """Add roman numeral item: (i) text"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(f"({numeral})  {text}")
    set_font(run)
    return p

def add_field_line(label, value=""):
    """Add form field line with label and blank/value."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.5)
    
    run1 = p.add_run(f"{label}: ")
    set_font(run1, bold=True)
    
    if value:
        run2 = p.add_run(value)
        set_font(run2)
    else:
        run2 = p.add_run("_" * 50)
        set_font(run2)
    return p

def add_line():
    """Add horizontal line separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 70)
    set_font(run)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_break():
    """Add paragraph break."""
    doc.add_paragraph()


# ============ DOCUMENT CONTENT ============

# TITLE
add_title("INDEPENDENT CONTRACTOR SERVICES AGREEMENT")
add_break()

add_centered("Agreement Reference: SA-2024-001")
add_centered("Effective Date: _______________________")
add_line()

# PARTIES
add_article("PARTIES")

add_plain("THE COMPANY", "THE COMPANY")
add_break()

add_field_line("Company Name")
add_field_line("SIRET / Registration No.")
add_field_line("Registered Address")
add_field_line("Country of Incorporation")
add_field_line("Represented By")
add_field_line("Position")
add_field_line("Email")
add_break()

add_para('Hereinafter referred to as "the Company" or "Client".', first_indent=0)
add_break()

add_plain("THE CONTRACTOR", "THE CONTRACTOR")
add_break()

add_field_line("Full Name", "Tomás Maria Burnay Batalha")
add_field_line("Portuguese NIF")
add_field_line("Address")
add_field_line("Country of Tax Residence", "Portugal")
add_field_line("Email")
add_field_line("Phone")
add_break()

add_para('Hereinafter referred to as "the Contractor" or "Service Provider".', first_indent=0)
add_line()

# RECITALS
add_article("RECITALS")

add_plain("WHEREAS:", "WHEREAS:")
add_break()

add_subitem("A", "The Company operates a healthtech business and seeks to expand its commercial operations through improved go-to-market capabilities;")
add_subitem("B", "The Contractor is an independent professional with expertise in commercial strategy, software development, and growth operations;")
add_subitem("C", "The Contractor operates as a self-employed professional (trabalhador independente) under Portuguese law and issues invoices via the Portuguese Tax Authority's Recibos Verdes system;")
add_subitem("D", "Both parties wish to establish a clear framework for the provision of services, compensation, and mutual obligations;")
add_break()

add_plain("NOW, THEREFORE, the parties agree as follows:", "NOW, THEREFORE,")
add_line()

# ARTICLE 1
add_article("ARTICLE 1: NATURE OF RELATIONSHIP")

add_section("1.1", "Independent Contractor Status", 
    "The Contractor is engaged as an independent contractor (prestataire de services / prestador de serviços) and not as an employee. There is no relationship of subordination (lien de subordination) between the parties.")

add_section("1.2", "No Employment Relationship", 
    "This Agreement does not create an employment contract (contrat de travail / contrato de trabalho). The Contractor:")

add_subitem("a", "Determines their own working hours and methods;")
add_subitem("b", "Provides their own equipment and workspace;")
add_subitem("c", "Is responsible for their own taxes, social security, and insurance;")
add_subitem("d", "May provide services to other clients.")

add_section("1.3", "EU Services Freedom", 
    "This Agreement is entered into under the freedom to provide services within the European Union, in accordance with Articles 56-62 of the Treaty on the Functioning of the European Union (TFEU).")

# ARTICLE 2
add_article("ARTICLE 2: SCOPE OF SERVICES")

add_section("2.1", "Primary Services", 
    "The Contractor agrees to provide the following services:")
add_break()

add_para("A. GTM System Development", first_indent=0.5)
add_subitem("i", "Custom Go-To-Market web application;")
add_subitem("ii", "AI-personalized bulk email campaign functionality;")
add_subitem("iii", "Lead import and management system;")
add_subitem("iv", "Sales pipeline visualization dashboard;")
add_subitem("v", "Email open/click tracking and analytics;")
add_subitem("vi", "Campaign scheduling and management interface.")
add_break()

add_para("B. Commercial Growth Services", first_indent=0.5)
add_subitem("i", "Outbound lead generation and prospecting;")
add_subitem("ii", "Sales campaign design and execution;")
add_subitem("iii", "Pipeline management and optimization;")
add_subitem("iv", "Performance reporting.")

add_section("2.2", "Minimum Viable Product (MVP)", 
    "The GTM system MVP shall include the core features listed in Section 2.1(A) in functional, deployable form suitable for business operations.")

add_section("2.3", "Time Commitment", 
    "The Contractor commits to approximately fifteen to twenty (15-20) hours per week during the term of this Agreement.")

add_section("2.4", "Exclusions", 
    "Unless separately agreed in writing, Services do not include: hardware procurement, third-party software licenses, legal or accounting advice, or in-person attendance at Company premises.")

# ARTICLE 3
add_article("ARTICLE 3: COMPENSATION")

add_section("3.1", "Advance Payment", 
    "Upon signing this Agreement, the Company shall pay:")

add_subitem("a", "Signing Advance: € _____________ (amount in words: _____________)")
add_subitem("b", "Tool & Infrastructure Budget: € 150.00")
add_subitem("c", "Total Upfront Payment: € _____________")

add_section("3.2", "Work Equivalent", 
    "The Signing Advance represents pre-payment for approximately sixty (60) hours of work at an effective rate of € _____________ per hour.")

add_section("3.3", "Monthly Compensation", 
    "Following the advance period:")

add_subitem("a", "Monthly Retainer: € _____________ per month")
add_subitem("b", "Commission on Closed Deals: _______% of net deal value")

add_section("3.4", "Payment Terms", "")

add_subitem("a", "Advance: Due within three (3) business days of signing;")
add_subitem("b", "Monthly Retainer: Due on the 1st of each calendar month;")
add_subitem("c", "Commission: Due within fifteen (15) days of Company receiving deal payment.")

add_section("3.5", "Running Costs", 
    "The Company shall reimburse ongoing operational costs for the GTM system, estimated at €50 per month, covering hosting, email infrastructure, and AI API usage.")

add_section("3.6", "Payment Method", 
    "All payments shall be made via bank transfer to the Contractor's Portuguese bank account:")

add_field_line("Account Holder", "Tomás Maria Burnay Batalha")
add_field_line("IBAN")
add_field_line("BIC/SWIFT")
add_field_line("Bank Name")

# ARTICLE 4
add_article("ARTICLE 4: INVOICING AND TAX COMPLIANCE")

add_section("4.1", "Invoicing", 
    "The Contractor shall issue invoices (Fatura-Recibo) via the Portuguese Tax Authority portal (Portal das Finanças) for all amounts due.")

add_section("4.2", "Invoice Requirements", 
    "Each invoice shall include:")

add_subitem("a", "Contractor's name, address, and Portuguese NIF;")
add_subitem("b", "Company's name, address, and SIRET/registration number;")
add_subitem("c", "Description of services rendered;")
add_subitem("d", "Amount in EUR;")
add_subitem("e", "VAT notation as per Section 4.3.")

add_section("4.3", "VAT Treatment", 
    "As an intra-EU B2B service:")

add_subitem("a", 'If Contractor is VAT-exempt (regime de isenção): Invoice states "IVA – regime de isenção, artigo 53.º do CIVA";')
add_subitem("b", 'If reverse charge applies: Invoice states "IVA – autoliquidação / Reverse charge – Article 196, EU VAT Directive 2006/112/EC".')

add_section("4.4", "Tax Responsibility", 
    "The Contractor is solely responsible for:")

add_subitem("a", "Portuguese income tax (IRS) obligations;")
add_subitem("b", "Portuguese social security (Segurança Social) contributions;")
add_subitem("c", "Any applicable VAT obligations.")

add_section("4.5", "No Withholding", 
    "The Company shall pay gross amounts without withholding, as no withholding tax applies to EU cross-border B2B services between Portugal and France.")

# ARTICLE 5
add_article("ARTICLE 5: DELIVERABLES AND TIMELINE")

add_section("5.1", "MVP Delivery Deadline", 
    "The Contractor shall deliver a functional GTM System MVP within six (6) weeks of receiving the Advance Payment.")

add_para("Advance Payment Received: _______________________ (Date)", indent=0.5, first_indent=0)
add_para("MVP Deadline: _______________________ (Date)", indent=0.5, first_indent=0)

add_section("5.2", "Progress Reporting", 
    "The Contractor shall provide weekly written updates including:")

add_subitem("a", "Hours worked during the period;")
add_subitem("b", "Tasks completed;")
add_subitem("c", "Tasks planned for the following week;")
add_subitem("d", "Blockers, risks, or concerns.")

add_section("5.3", "Acceptance Process", 
    "The Company shall have seven (7) calendar days following MVP delivery to:")

add_subitem("a", "Review the deliverables;")
add_subitem("b", "Request reasonable modifications;")
add_subitem("c", "Provide written acceptance or objections.")

add_para("If no written response is received within seven (7) days, the MVP shall be deemed accepted.", first_indent=0.5)

# ARTICLE 6
add_article("ARTICLE 6: ADVANCE REFUND CONDITIONS")

add_section("6.1", "Refund Triggers (Contractor Responsibility)", 
    "The Contractor shall refund a proportionate amount of the Advance if:")

add_subitem("a", "Contractor ceases communication for fourteen (14) or more consecutive days without valid reason;")
add_subitem("b", "Contractor explicitly terminates before MVP delivery without mutual consent;")
add_subitem("c", "Contractor fails to deliver MVP by deadline (excluding Company-caused delays).")

add_section("6.2", "Refund Calculation", 
    "The refund amount shall be calculated as follows:")

add_centered("Refund Amount = Advance Payment − (Hours Worked × Hourly Rate)", italic=True)

add_para("Hours Worked shall be determined by the weekly progress reports under Section 5.2.", first_indent=0.5)

add_section("6.3", "No Refund Conditions", 
    "The Contractor retains the full Advance with no refund obligation if:")

add_subitem("a", "MVP is delivered as specified;")
add_subitem("b", "Company terminates for reasons other than Contractor's material breach;")
add_subitem("c", "Company fails to respond for fourteen (14) or more consecutive days;")
add_subitem("d", "Company materially changes scope without written agreement;")
add_subitem("e", "Mutual agreement to end after equivalent hours worked;")
add_subitem("f", "Force majeure, illness, or emergency (with prompt communication).")

add_section("6.4", "Refund Timeline", 
    "Any refund due shall be paid within thirty (30) calendar days of the triggering event.")

# ARTICLE 7
add_article("ARTICLE 7: INTELLECTUAL PROPERTY")

add_section("7.1", "Assignment", 
    "Upon full payment of all compensation due under this Agreement, all intellectual property rights in the GTM system and deliverables created under this Agreement shall transfer to the Company.")

add_section("7.2", "Pre-Existing IP", 
    "The Contractor retains all rights to pre-existing tools, frameworks, libraries, and methodologies. The Company receives a perpetual, non-exclusive license to use any pre-existing IP incorporated in the deliverables.")

add_section("7.3", "Interim License", 
    "Until full payment is received, the Company receives a limited, revocable license to use the deliverables for internal business purposes only.")

# ARTICLE 8
add_article("ARTICLE 8: CONFIDENTIALITY")

add_section("8.1", "Confidential Information", 
    "Both parties agree to maintain the confidentiality of proprietary information disclosed during the engagement, including: business strategies, customer data, technical specifications, pricing, and financial information.")

add_section("8.2", "Duration", 
    "Confidentiality obligations survive termination for two (2) years.")

add_section("8.3", "Permitted Disclosures", 
    "Confidentiality obligations do not apply to information that:")

add_subitem("a", "Becomes publicly available through no fault of the receiving party;")
add_subitem("b", "Was already known to the receiving party prior to disclosure;")
add_subitem("c", "Is independently developed by the receiving party;")
add_subitem("d", "Is required to be disclosed by law or court order.")

# ARTICLE 9
add_article("ARTICLE 9: GOOD FAITH AND DISPUTE RESOLUTION")

add_section("9.1", "Good Faith", 
    "Both parties commit to acting in good faith and communicating openly about any concerns or issues.")

add_section("9.2", "Notice Before Breach Claim", 
    "Neither party may claim material breach without first providing seven (7) days written notice specifying:")

add_subitem("a", "The alleged breach;")
add_subitem("b", "The remedy requested;")
add_subitem("c", "A reasonable opportunity to cure.")

add_section("9.3", "Escalation Process", 
    "The parties agree to the following dispute resolution process:")

add_subitem("a", "Step 1 — Direct Discussion (14 days): Parties attempt to resolve the dispute directly;")
add_subitem("b", "Step 2 — Mediation (30 days): If unresolved, neutral third-party mediation;")
add_subitem("c", "Step 3 — Legal Action: If mediation fails, court proceedings may be initiated.")

add_section("9.4", "Jurisdiction", 
    "Any legal proceedings shall be brought before the courts of (select one by mutual agreement):")

add_subitem("a", "☐ Paris, France;")
add_subitem("b", "☐ Lisbon, Portugal;")
add_subitem("c", "☐ Other: _________________________________.")

# ARTICLE 10
add_article("ARTICLE 10: TERM AND TERMINATION")

add_section("10.1", "Initial Term", 
    "This Agreement begins on the Effective Date and continues for three (3) months.")

add_section("10.2", "Renewal", 
    "After the initial term, this Agreement renews month-to-month unless either party provides fourteen (14) days written notice of non-renewal.")

add_section("10.3", "Termination for Convenience", 
    "Either party may terminate with fourteen (14) days written notice, subject to the refund provisions in Article 6.")

add_section("10.4", "Termination for Cause", 
    "Either party may terminate immediately upon material breach, following the notice procedure in Section 9.2.")

add_section("10.5", "Effect of Termination", 
    "Upon termination:")

add_subitem("a", "All unpaid amounts become immediately due;")
add_subitem("b", "Confidentiality obligations continue per Section 8.2;")
add_subitem("c", "IP provisions apply per Article 7.")

# ARTICLE 11
add_article("ARTICLE 11: GENERAL PROVISIONS")

add_section("11.1", "Language", 
    "This Agreement is drafted in English by mutual consent of the parties.")

add_section("11.2", "Governing Law", 
    "This Agreement shall be governed by and construed in accordance with (select one by mutual agreement):")

add_subitem("a", "☐ French law (droit français);")
add_subitem("b", "☐ Portuguese law (lei portuguesa);")
add_subitem("c", "☐ Other: _________________________________.")

add_section("11.3", "Entire Agreement", 
    "This Agreement constitutes the entire agreement between the parties and supersedes all prior negotiations, representations, and agreements.")

add_section("11.4", "Amendments", 
    "Any amendments must be in writing and signed by both parties.")

add_section("11.5", "Severability", 
    "If any provision is found unenforceable, the remaining provisions continue in full effect.")

add_section("11.6", "Notices", 
    "All notices shall be in writing and sent via email to the addresses specified in the Parties section, with read receipt requested.")

add_section("11.7", "Counterparts", 
    "This Agreement may be executed in counterparts, including electronic signatures, each of which shall constitute an original.")

add_line()

# SIGNATURES
add_article("SIGNATURES")

add_para("By signing below, both parties confirm they have read, understood, and agree to be bound by the terms of this Agreement.", first_indent=0)
add_break()
add_break()

add_plain("THE COMPANY", "THE COMPANY")
add_break()

add_para("Signature: _________________________________________________", first_indent=0)
add_para("Printed Name: _________________________________________________", first_indent=0)
add_para("Title: _________________________________________________", first_indent=0)
add_para("Date: _________________________________________________", first_indent=0)
add_para("Place: _________________________________________________", first_indent=0)
add_break()
add_break()

add_plain("THE CONTRACTOR", "THE CONTRACTOR")
add_break()

add_para("Signature: _________________________________________________", first_indent=0)
add_para("Printed Name: Tomás Maria Burnay Batalha", first_indent=0)
add_para("Date: _________________________________________________", first_indent=0)
add_para("Place: Portugal", first_indent=0)

add_line()

# EXHIBIT A
add_title("EXHIBIT A: PAYMENT SCHEDULE")
add_break()

add_para("The following payments are due under this Agreement:", first_indent=0)
add_break()

add_para("1. Signing Advance: € _________ — Due upon signing — Status: ☐ Paid", first_indent=0.5)
add_para("2. Tool Budget: € 150.00 — Due upon signing — Status: ☐ Paid", first_indent=0.5)
add_para("3. Month 1 Retainer: € _________ — Due: _____________ — Status: ☐ Paid", first_indent=0.5)
add_para("4. Month 2 Retainer: € _________ — Due: _____________ — Status: ☐ Paid", first_indent=0.5)
add_para("5. Month 3 Retainer: € _________ — Due: _____________ — Status: ☐ Paid", first_indent=0.5)
add_para("6. Running Costs: ~€ 50.00 per month — Due: 1st of each month — Status: ☐ Paid", first_indent=0.5)

add_break()

# EXHIBIT B
add_title("EXHIBIT B: GTM SYSTEM MVP SPECIFICATIONS")
add_break()

add_para("The Minimum Viable Product shall include the following features:", first_indent=0)
add_break()

add_para("1. Lead Management (Must Have): Import CSV/Excel, search, filter, tag leads.", first_indent=0.5)
add_para("2. Email Campaigns (Must Have): Create, schedule, send bulk personalized emails.", first_indent=0.5)
add_para("3. AI Personalization (Must Have): GPT-powered email customization per lead.", first_indent=0.5)
add_para("4. Pipeline Dashboard (Must Have): Visual Kanban board for sales stages.", first_indent=0.5)
add_para("5. Email Tracking (Must Have): Open and click tracking with analytics.", first_indent=0.5)
add_para("6. User Authentication (Must Have): Secure login for Company team members.", first_indent=0.5)
add_para("7. Basic Reporting (Should Have): Campaign performance metrics.", first_indent=0.5)

add_line()

# END
add_centered("— END OF AGREEMENT —", size=12, bold=True)
add_break()
add_centered("Document Reference: SA-2024-001 | Prepared December 2024", size=9, italic=True)

# Save
output_path = '/Users/tomasbatalha/Downloads/Tomas_Batalha_Future_Plan/1_Current_Priorities/Korean_Startup/Services_Agreement_FINAL.docx'
doc.save(output_path)
print(f'✅ Created: {output_path}')
print('')
print('Standard Legal Format Applied:')
print('  • No tables - replaced with indented text blocks')
print('  • No bullet points - using (a), (b), (c) enumeration')
print('  • Hierarchical numbering: Article → 1.1 → (a)')
print('  • Times New Roman throughout')
print('  • 1.25" left margin, 1" other margins')
print('  • Justified alignment')
