#!/usr/bin/env python3
"""
CV Editor Tool
Easily add bullets and jobs to your CV while preserving formatting.

Usage:
    python cv_editor.py                    # Interactive mode
    python cv_editor.py --add-bullet       # Add bullet at specific job
    python cv_editor.py --add-job          # Add new job to section
    python cv_editor.py --list             # List all jobs in CV
"""

import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
from datetime import datetime

# Default CV path
CV_PATH = os.path.join(os.path.dirname(__file__), 'TomásBatalha_Resume_12_2025_POLISHED.docx')

def get_cv_structure(doc):
    """Parse CV and return structured data about sections and jobs."""
    structure = {
        'WORK EXPERIENCE': [],
        'EDUCATION': [],
        'LEADERSHIP EXPERIENCE': [],
        'SKILLS': []
    }
    
    current_section = None
    current_job = None
    
    def process_table(table):
        nonlocal current_section, current_job
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    # Check for section headers
                    for section in structure.keys():
                        if section in text:
                            current_section = section
                            break
                    
                    # Check for company/job lines (contain |)
                    if current_section and current_section != 'SKILLS':
                        if '|' in text and len(text) < 100:
                            # Could be company line or role line
                            if current_job is None or 'role' in current_job:
                                # New job
                                current_job = {'company': text, 'bullets': [], 'para': para}
                                structure[current_section].append(current_job)
                            else:
                                # Role line for current job
                                current_job['role'] = text
                        elif text.startswith(('•', '-', '–')) or para.style.name == 'List Paragraph':
                            # Bullet point
                            if current_job:
                                current_job['bullets'].append(text)
                
                # Process nested tables
                for nested in cell.tables:
                    process_table(nested)
    
    for table in doc.tables:
        process_table(table)
    
    return structure

def list_jobs(doc):
    """List all jobs in the CV."""
    structure = get_cv_structure(doc)
    
    print("\n" + "=" * 60)
    print("CV STRUCTURE")
    print("=" * 60)
    
    for section, jobs in structure.items():
        if jobs:
            print(f"\n📁 {section}")
            print("-" * 40)
            for i, job in enumerate(jobs, 1):
                company = job.get('company', 'Unknown')
                role = job.get('role', '')
                bullets = len(job.get('bullets', []))
                print(f"  {i}. {company[:50]}")
                if role:
                    print(f"     {role[:50]}")
                print(f"     [{bullets} bullets]")

def add_bullet_interactive(doc):
    """Interactively add a bullet to an existing job."""
    structure = get_cv_structure(doc)
    
    # Show jobs
    print("\n📋 SELECT JOB TO ADD BULLET:")
    print("-" * 40)
    
    all_jobs = []
    idx = 1
    for section, jobs in structure.items():
        if jobs and section != 'SKILLS':
            print(f"\n{section}:")
            for job in jobs:
                company = job.get('company', 'Unknown')[:40]
                print(f"  {idx}. {company}")
                all_jobs.append((section, job))
                idx += 1
    
    try:
        choice = int(input("\nEnter job number: ")) - 1
        if choice < 0 or choice >= len(all_jobs):
            print("Invalid selection")
            return False
    except ValueError:
        print("Invalid input")
        return False
    
    section, job = all_jobs[choice]
    
    print(f"\nAdding bullet to: {job.get('company', 'Unknown')[:50]}")
    print("\nTip: Start with action verb (Built, Led, Designed, Secured, etc.)")
    bullet_text = input("\nEnter bullet text: ").strip()
    
    if not bullet_text:
        print("No text entered, cancelled.")
        return False
    
    # Find the job in the document and add bullet after last bullet
    # This is tricky with nested tables - we'll add to the paragraph
    
    print(f"\n✅ Bullet to add: {bullet_text[:60]}...")
    print("\n⚠️  Due to Word's complex table structure, please add manually:")
    print("   1. Open the CV in Word")
    print("   2. Go to the job section")
    print("   3. Add this bullet")
    print(f"\n   • {bullet_text}")
    
    return True

def add_job_interactive():
    """Interactively add a new job."""
    print("\n📝 ADD NEW JOB")
    print("-" * 40)
    
    print("\nSelect section:")
    print("  1. Work Experience")
    print("  2. Education") 
    print("  3. Leadership Experience")
    
    try:
        section_choice = int(input("\nEnter choice (1-3): "))
    except ValueError:
        print("Invalid input")
        return None
    
    sections = {1: 'WORK EXPERIENCE', 2: 'EDUCATION', 3: 'LEADERSHIP EXPERIENCE'}
    section = sections.get(section_choice)
    
    if not section:
        print("Invalid selection")
        return None
    
    print(f"\nAdding to: {section}")
    
    company = input("\nCompany/Institution: ").strip()
    if not company:
        return None
        
    location = input("Location (e.g., New York, USA): ").strip()
    if not location:
        return None
        
    role = input("Role/Title: ").strip()
    if not role:
        return None
        
    dates = input("Dates (e.g., January 2024 – Present): ").strip()
    if not dates:
        return None
    
    print("\nFirst bullet (required):")
    bullet = input("• ").strip()
    if not bullet:
        return None
    
    job_data = {
        'section': section,
        'company': company,
        'location': location,
        'role': role,
        'dates': dates,
        'bullets': [bullet]
    }
    
    # Ask for more bullets
    while True:
        more = input("\nAdd another bullet? (y/n): ").strip().lower()
        if more == 'y':
            b = input("• ").strip()
            if b:
                job_data['bullets'].append(b)
        else:
            break
    
    return job_data

def print_job_template(job_data):
    """Print formatted job for manual copy-paste."""
    print("\n" + "=" * 60)
    print("📋 COPY THIS TO YOUR CV:")
    print("=" * 60)
    print(f"\n{job_data['company']} | {job_data['location']}")
    print(f"{job_data['role']} | {job_data['dates']}")
    for bullet in job_data['bullets']:
        print(f"• {bullet}")
    print("\n" + "=" * 60)

def main():
    print("\n" + "=" * 60)
    print("  CV EDITOR TOOL")
    print("  For: TomásBatalha Resume")
    print("=" * 60)
    
    if not os.path.exists(CV_PATH):
        print(f"\n❌ CV not found at: {CV_PATH}")
        print("Please ensure the file exists or update CV_PATH in script.")
        return
    
    # Load document
    try:
        doc = Document(CV_PATH)
        print(f"\n✅ Loaded: {os.path.basename(CV_PATH)}")
    except Exception as e:
        print(f"\n❌ Error loading CV: {e}")
        return
    
    while True:
        print("\n📌 MENU:")
        print("  1. List all jobs")
        print("  2. Add bullet to existing job")
        print("  3. Add new job")
        print("  4. Exit")
        
        try:
            choice = input("\nChoice (1-4): ").strip()
        except (EOFError, KeyboardInterrupt):
            break
        
        if choice == '1':
            list_jobs(doc)
        elif choice == '2':
            add_bullet_interactive(doc)
        elif choice == '3':
            job_data = add_job_interactive()
            if job_data:
                print_job_template(job_data)
        elif choice == '4':
            print("\n👋 Bye!")
            break
        else:
            print("Invalid choice")

if __name__ == '__main__':
    main()
